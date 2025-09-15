import os
import uuid
from pathlib import Path

MEDIA_ROOT = Path(__file__).resolve().parent.parent / 'media'
CHROMA_DIR = Path(__file__).resolve().parent.parent / 'chroma_db'
os.makedirs(CHROMA_DIR, exist_ok=True)

def extract_text_from_file(filepath: str) -> str:
    filepath = str(filepath)
    lower = filepath.lower()
    text = ''
    if lower.endswith('.pdf'):
        try:
            from pypdf import PdfReader
            reader = PdfReader(filepath)
            pages = [p.extract_text() or '' for p in reader.pages]
            text = '\n'.join(pages)
        except Exception as e:
            text = ''
    elif lower.endswith('.docx'):
        try:
            import docx
            doc = docx.Document(filepath)
            paragraphs = [p.text for p in doc.paragraphs]
            text = '\n'.join(paragraphs)
        except Exception as e:
            text = ''
    else:
        # fallback for txt or others
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        except Exception:
            text = ''
    return text or ''

def chunk_text(text: str, chunk_size: int = 800, overlap: int = 100):
    if not text:
        return []
    chunks = []
    start = 0
    length = len(text)
    while start < length:
        end = start + chunk_size
        chunk = text[start:end]
        chunks.append(chunk)
        start = end - overlap
        if start < 0:
            start = 0
    return chunks

# Embedding & vector store helpers (lazy imports)
def get_embedding_model(model_name: str = None):
    model_name = model_name or os.environ.get('EMBEDDING_MODEL', 'all-MiniLM-L6-v2')
    try:
        from sentence_transformers import SentenceTransformer
        emb = SentenceTransformer(model_name)
        return emb
    except Exception as e:
        raise RuntimeError(f"Failed to load embedding model: {e}") from e

def get_chroma_client_and_collection(collection_name: str = 'documents'):
    try:
        import chromadb
        from chromadb.config import Settings
        client = chromadb.Client(Settings(chroma_db_impl="duckdb+parquet", persist_directory=str(CHROMA_DIR)))
        # create or get collection
        coll = None
        try:
            coll = client.get_collection(name=collection_name)
        except Exception:
            coll = client.create_collection(name=collection_name)
        return client, coll
    except Exception as e:
        raise RuntimeError(f"Failed to initialize ChromaDB: {e}") from e

def embed_texts(texts, embedder=None):
    # returns list of embeddings (list of floats)
    if embedder is None:
        embedder = get_embedding_model()
    return embedder.encode(texts, show_progress_bar=False, convert_to_numpy=True).tolist()

# chatapp/utils.py
from pypdf import PdfReader
import docx
from typing import List

def extract_text_from_file(path: str) -> List[str]:
    if path.lower().endswith(".pdf"):
        try:
            reader = PdfReader(path)
            pages = []
            for p in reader.pages:
                text = p.extract_text() or ""
                if text.strip():
                    pages.append(text)
            return pages
        except Exception as e:
            return []
    elif path.lower().endswith(".docx"):
        doc = docx.Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return paragraphs
    else:
        # try plain text
        try:
            with open(path, "r", encoding="utf-8") as fh:
                return [fh.read()]
        except:
            return []

# chatapp/utils.py (continuing)
from sentence_transformers import SentenceTransformer
import chromadb
from chromadb.config import Settings
from chromadb.utils import embedding_functions
import os

CHROMA_DIR = os.getenv("CHROMA_DIR", "chroma_db")

# init once (module-level)
client = chromadb.Client(Settings(chroma_db_impl="duckdb+parquet", persist_directory=CHROMA_DIR))
embedding_model_name = os.getenv("EMBEDDING_MODEL", "all-MiniLM-L6-v2")
st_model = SentenceTransformer(embedding_model_name)
ef = embedding_functions.SentenceTransformerEmbeddingFunction(model_name=embedding_model_name, sentence_transformer=st_model)

def add_documents_to_vector_store(text_chunks, doc_id="doc"):
    col_name = "documents"
    # create or get collection
    collection = client.get_or_create_collection(name=col_name, embedding_function=ef)
    ids = [f"{doc_id}_{i}" for i in range(len(text_chunks))]
    metadatas = [{"source": doc_id, "chunk": i} for i in range(len(text_chunks))]
    collection.add(
        documents=text_chunks,
        metadatas=metadatas,
        ids=ids
    )
    client.persist()

    # utils.py
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
import torch

MODEL_NAME = "google/flan-t5-small"  # small & CPU-friendly
device = "cuda" if torch.cuda.is_available() else "cpu"

tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME)
model = AutoModelForSeq2SeqLM.from_pretrained(MODEL_NAME).to(device)

def generate_answer(question: str, contexts: list):
    context_text = "\n\n".join(contexts)
    prompt = f"Answer the question based on the context. If not found, say 'I don't know'.\n\nContext:\n{context_text}\n\nQuestion: {question}\nAnswer:"
    inputs = tokenizer(prompt, return_tensors="pt", truncation=True, max_length=512).to(device)
    outputs = model.generate(**inputs, max_new_tokens=200)
    return tokenizer.decode(outputs[0], skip_special_tokens=True)

from llama_cpp import Llama

llm = Llama(model_path="models/mistral-7b-instruct-v0.1.Q4_K_M.gguf", n_ctx=2048)

def generate_answer(question, contexts):
    context_text = "\n\n".join(contexts)
    prompt = f"Context:\n{context_text}\n\nQuestion: {question}\nAnswer:"
    output = llm(prompt, max_tokens=256)
    return output["choices"][0]["text"].strip()
